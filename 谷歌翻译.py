# -*- coding:utf-8 -*
import execjs
from urllib.request import quote
import requests
#import pymysql
import os
import re
import docx

'''---谷歌翻译软件---'''
class Py4Js():
    """获取加密参数tk"""
    def __init__(self):
        self.ctx = execjs.compile("""
        function TL(a) {
        var k = "";
        var b = 406644;
        var b1 = 3293161072;
        var jd = ".";
        var $b = "+-a^+6";
        var Zb = "+-3^+b+-f";
        for (var e = [], f = 0, g = 0; g < a.length; g++) {
            var m = a.charCodeAt(g);
            128 > m ? e[f++] = m : (2048 > m ? e[f++] = m >> 6 | 192 : (55296 == (m & 64512) && g + 1 < a.length && 56320 == (a.charCodeAt(g + 1) & 64512) ? (m = 65536 + ((m & 1023) << 10) + (a.charCodeAt(++g) & 1023),
            e[f++] = m >> 18 | 240,
            e[f++] = m >> 12 & 63 | 128) : e[f++] = m >> 12 | 224,
            e[f++] = m >> 6 & 63 | 128),
            e[f++] = m & 63 | 128)
        }
        a = b;
        for (f = 0; f < e.length; f++) a += e[f],
        a = RL(a, $b);
        a = RL(a, Zb);
        a ^= b1 || 0;
        0 > a && (a = (a & 2147483647) + 2147483648);
        a %= 1E6;
        return a.toString() + jd + (a ^ b)
    };
    function RL(a, b) {
        var t = "a";
        var Yb = "+";
        for (var c = 0; c < b.length - 2; c += 3) {
            var d = b.charAt(c + 2),
            d = d >= t ? d.charCodeAt(0) - 87 : Number(d),
            d = b.charAt(c + 1) == Yb ? a >>> d: a << d;
            a = b.charAt(c) == Yb ? a + d & 4294967295 : a ^ d
        }
        return a
    }
    """)

    def getTk(self, text):
        return self.ctx.call("TL", text)

def open_url(url):
    """
    获取原始文本
    注意：google会对每个ip进行检测，每个ip只能调用上千次，建议请求时带上代理
    :param url:
    :return:
    """
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64; rv:23.0) Gecko/20100101 Firefox/23.0'}
    data = requests.get(url=url, headers=headers).text
    return data
def translate(content,transway):
    """
    翻译，返回翻译的结果
    注意：本例是英译中，要是中译英要把url中的‘sl=en&tl=zh-CN’换成‘sl=zh-CN&tl=en’即可
    :param content:传入待翻译的文本
    :return:
    """
    js = Py4Js()
    tk = js.getTk(content)
    content = quote(content)
    url = "https://translate.google.cn/translate_a/single?client=webapp&"+transway+"&hl=EN&dt=at&dt=bd&dt=ex&dt=ld&dt=md&dt=qca" \
          "&dt=rw&dt=rm&dt=ss&dt=t&ie=UTF-8&oe=UTF-8&clearbtn=1&otf=1&pc=1" \
          "&srcrom=0&ssel=0&tsel=0&kc=2&tk=%s&q=%s" % (tk, content)

    result = open_url(url)
    end = result.find("\",")
    if end > 4:
        texts = result[4:end]
        return texts
'''---换头工具---'''
#def check(document,timesworld_break):
    #for num_countuseless in range(0,timesworld_break):
        #fisrtone = input("想要被替换的第"+str(num_countuseless+1)+"个字符串")
        #secondone=input("想要替换的新字符串")
        #DICT = {
            #fisrtone: secondone
        #}
    # tables
        #for table in document.tables:
           #for row in range(len(table.rows)):
                #for col in range(len(table.columns)):
                   #for key, value in DICT.items():
                       #if key in table.cell(row, col).text:
                            #print(key + "->" + value)
                            #table.cell(row, col).text = table.cell(row, col).text.replace(key, value)

    # paragraphs
        #for para in document.paragraphs:
            #for i in range(len(para.runs)):
                #for key, value in DICT.items():
                    #if key in para.runs[i].text:
                        #print(key + "->" + value)
                        #para.runs[i].text = para.runs[i].text.replace(key, value)
    #print("替换完成")

    #return document

#def worldbreak():
    #PATH = r"./docx库/"
    #oldname = input("请确认您的文件的具体位置（“翻译库/”；“跑团log库/”；“新闻库/”；“/”）以及文件名:")
    #oldFile = PATH + oldname +".docx"
    #newname=input("您的文件需要一个新名字")
    #times_worldbreak=eval(input("您要更替的次数"))
    #newFile = PATH + "翻译库/" + newname + ".docx"
    #document = Document(oldFile)
    #document = check(document,times_worldbreak)
    #document.save(newFile)
    #return newname
'''----翻译很长的话----'''
def diaoyongdocx(transway):
    if eval(transway) == 1:
        transway = "sl=en&tl=zh-CN"
        print('翻译模式为英译中')
    else:
        transway = "sl=zh-CN&tl=en"
        print('翻译模式为中译英')
    doc = docx.Document()
    doc.save(r'./待翻译文段.docx')
    os.system('待翻译文段.docx')
    doc = docx.Document(r'./待翻译文段.docx')
    new_doc = docx.Document()
    # 遍历每一段文本
    for para in doc.paragraphs:
        print(para.text)
        pattern = r'\.|/|;|`|\[|\]|<|>|\?|\{|\}|\~|!|@|#|\$|%|\^|&|\(|\)|-|=|\_|\+|。|、|；|‘|’|【|】|！|（|）'
        result_list = re.split(pattern, para.text)
        len = ''
        for words in result_list:
            trans = translate(words, transway)
            print(trans)
            len += trans
            len += "。"
            print(len.replace('l,null,"en。',''))
        new_doc.add_paragraph(para.text)
        new_doc.add_paragraph(len.replace('l,null,"en。',''))
    new_doc.save(r"./翻译文段.docx")
    os.system('翻译文段.docx')


'''----鬼畜命令----'''
def guichu():
    transway = "sl=en&tl=zh-CN"
    doc = input('输入您要翻译的语句或命令:')
    if doc == '我要翻译很长的话':
        doc = '打开doc文件'
        transway = input("请选择翻译方式,1是English翻译中文，2是中文翻译为English,默认为：")
        return doc,transway
    if doc == 'end':
        return doc,transway
    else:
        if doc == '打开记事本':
            os.system('notepad')
            (doc,transway) = guichu()
            return doc,transway
    print("语句已输入")
    try:
        transway = input("请选择翻译方式,1是English翻译中文，2是中文翻译为English,默认为：")
        if eval(transway) == 1:
            transway = "sl=en&tl=zh-CN"
        else:
            transway = "sl=zh-CN&tl=en"
    except:
        transway = "sl=en&tl=zh-CN"
    return doc,transway
'''----主程序----'''
#newname = worldbreak()
keepon = 1
trun_notepad = '是'
while(keepon):
    #if turn_notepad == '是':
        #doc = input('待翻译文字是否需要先打开记事本？（回答是或否）：')
        #if doc == '是':
            #os.system('notepad')
        #else:
            #turn_notepad = input('好的，之后这个问题还要出现吗？')
            #print('好的，我记住了。')
    #else:
        #print('整理翻译删除换行用的记事本已关闭。想打开请在翻译语句中输入“打开记事本”')
    print('''----------------------翻译程序运行中---------------------------
在doc中输入“end”即可退出此程序
在doc中输入“打开记事本”即可打开记事本并对翻译文段进行编辑。
在doc中输入"我要翻译很长的话"来翻译很长的话。
翻译文段必须为完整的一段话，不可换行。
----------------------翻译程序运行中---------------------------''')
    (doc,transway) = guichu()
    if doc == '打开doc文件':
        diaoyongdocx(transway)
    if doc == 'end':
        break
#new_doc = docx.Document()
# 遍历每一段文本
#for para in doc:
    #print(para.text)
    if doc == '打开doc文件':
        print('翻译任务继续')
    else:
        pattern = r'\.|/|;|`|\[|\]|<|>|\?|\{|\}|\~|!|@|#|\$|%|\^|&|\(|\)|-|=|\_|\+|。|、|；|‘|’|【|】|！|（|）'
        result_list = re.split(pattern,doc)
        len = ''
        for words in result_list:
            trans = translate(words,transway)
            len += trans
            len += "。"
        len = len.replace('l,null,"en。','')
        print('-------------------翻译对象是--------------------------\n'+doc+'\n---------------------------翻译对象是-------------------------')
        print('-------------------翻译结果是--------------------------\n'+len+'\n---------------------------翻译结果是-------------------------')

